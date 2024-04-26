from docx import Document
import fitz
from pypdf import PdfReader, PdfWriter
import bs4 as BeautifulSoup
import requests
import io
import unicodedata
from syllabusSequencer import sequencer
from reqAndAuth import canvas

#dependencies
canvasClient = canvas()

class docFile():
    def parseDoc(self):
        result = []
        for para in self.content.paragraphs:
            result.append(stylizedTextDoc(para))
        return result
    def startSequencer(self):
        return(sequencer(self.markedTerms,self.text))
    def __init__(self, canvFile):
        response = requests.get(canvFile.url, canvasClient.header)
        source_stream = io.BytesIO(response.content)
        self.content= Document(source_stream)
        self.text = [para.text.replace("\t","") for para in self.content.paragraphs]
        self.sents = [sent for sent in nlp(" ".join(self.text)).sents]
        self.docParse= self.parseDoc()
        self.markedTerms = [item.markedTerms[0] for item in self.docParse if len(item.markedTerms) > 0]
        self.sequences = self.startSequencer()

class pdfFile():
    def parseHTML(self):
        result = []
        for para in self.html.find_all('p'):
            result.append(stylizedTextPDF(para, para.attrs['style']))     
        return result
    def startSequencer(self):
        return(sequencer(self.markedTerms,self.text))
    def __init__(self, canvFile):
        response = requests.get(canvFile.url, canvasClient.header)
        source_stream = io.BytesIO(response.content)
        try:
            self.content = fitz.open("pdf",stream=source_stream)
        except:
            self.html=None
            self.text="CORRUPTED FILE"
            self.sections =[]
            return
        self.html = BeautifulSoup("".join([page.get_text("html") for page in self.content]), 'html.parser')
        self.htmlParse =self.parseHTML()
        self.markedTerms = [item.markedTerms[0] for item in self.htmlParse if len(item.markedTerms) > 0]
        self.text =[unicodedata.normalize("NFKD", item.text).strip() for item in self.htmlParse]
        self.sents = [sent for sent in nlp(" ".join(self.text)).sents]
        self.sequences = self.startSequencer()

class stylizedTextDoc():
    def markTerms(self, paraObject):
        for run in paraObject.runs:
            if run.font.bold:
                self.markedTerms.append(self.text)
                self.containsBold = True
                continue
            if run.font.italic:
                self.markedTerms.append(self.text)
                self.containsItalics = True
                continue
            if run.font.underline:
                self.markedTerms.append(self.text)
                self.containsUnderline = True
                continue
            if run.text.isupper():
                self.markedTerms.append(self.text)
                self.containsAllCaps = True
                continue
    def __init__(self, paraObject):
        self.text = paraObject.text
        self.style = paraObject.style.name
        self.containsAllCaps = paraObject.text.isupper()
        self.containsBold = paraObject.style.font.bold
        self.containsItalics= paraObject.style.font.italic
        self.containsUnderline = paraObject.style.font.underline
        self.markedTerms=[]
        self.markTerms(paraObject)


class stylizedTextPDF():
    def markTerms(self):
        if self.containsBold:
            if hasattr(self.content.b, "text"):
                self.markedTerms.append(self.content.b.text)
        if self.containsItalics:
            if hasattr(self.content.i, "text"):
                self.markedTerms.append(self.content.i.text)
        if self.containsUnderline:
            if hasattr(self.content.u, "text"):
                self.markedTerms.append(self.content.u.text)
        if self.containsAllCaps:
            self.markedTerms.append(self.content.text)
    def __init__(self, paraObject, styleString):
        self.content = paraObject
        self.style = styleString
        self.text = paraObject.text
        self.markedTerms=[]
        self.containsAllCaps = self.text.isupper()
        self.containsBold = hasattr(self.content, "b")
        self.containsItalics = hasattr(self.content, "i")
        self.containsUnderline = hasattr(self.content, "u")
        self.markedTerms = []
        self.markTerms()








